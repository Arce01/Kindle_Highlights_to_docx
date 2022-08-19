# -*- coding: utf-8 -*-
"""author: Daniel Arce Carnero
   LinkedIn: https://www.linkedin.com/in/daniel-arce-carnero-31a014140
"""

from tkinter import filedialog
from bs4 import BeautifulSoup

import os
from os import replace
import re

from docx import Document

# Select the URL file
fichero = filedialog.askopenfilename(title='Abrir', filetypes= (('Ficheros de subrayado','*.html'),
('Todos los ficheros', '*.*')))


# Read the document
html_file = open(fichero, 'r')
all = html_file.read()
S = BeautifulSoup(all, 'lxml')
content = S.get_text()
html_file.close()

content1 = content

# Get rid of non content lines
content1 = re.sub("Location [0-9]+", "", content1, flags = re.IGNORECASE)

content1 = content1.replace("Highlight (yellow) -  ", "")
content1 = content1.replace("", "")


title = fichero[(fichero.index('Downloads')+len('Downloads/')):fichero.index(' - Note')]
content1 = content1.replace("\n\n\n", "\n")

notes = len(list(re.findall(r"\bNote\b", content1)))
for i in range(10):
  note_i = content1.find('Note -')
  note_i =+ 7

c = content1.splitlines()
c.append('Nothing')

for item in ['','Note -  ']:
  try:
    while c.index(item) > -1:
      c.remove(item)
  except:
    pass

#Create the document
doc = Document()

#Save the different parts of the docx as different types
doc.add_heading(title, 0)

for line in range(len(c)):
  try:
    #Title 1
    if c[line+1] == 'T1':
      doc.add_paragraph('')
      doc.add_heading(c[line].title(), 1)
      c.pop(line+1)
      
    #Title 2
    elif c[line+1] == 'T2':
      doc.add_paragraph('')
      doc.add_heading(c[line].title(), 2)
      c.pop(line+1)
    
    #Title 3
    elif c[line+1] == 'T3':
      doc.add_paragraph('')
      doc.add_heading(c[line].title(), 3)
      c.pop(line+1)
    
    #Title 4
    elif c[line+1] == 'T4':
      doc.add_paragraph('')
      doc.add_heading(c[line].capitalize(), 4)
      c.pop(line+1)
    
    #First bullet point
    elif c[line+1] == 'B1':
      
      fr = c[line:]
      idx2 = line + fr.index('B2') #Last bullet point

      B_items = c[line:idx2]
      B_items.remove('B1')

      for item in B_items:
        doc.add_paragraph(item.capitalize(), style = "List Bullet")
      
      doc.add_paragraph('')

      for idx in range(line,idx2):
        c.pop(line+1)
    
    #Literal text
    elif c[line+1] == 'L':
      texto = ' "' + c[line] + '"'

      p = doc.add_paragraph('')
      p.add_run(texto).italic = True
      c.pop(line+1)
    
    #The rest
    else:
      doc.add_paragraph(c[line].capitalize()) 
  
  except:
    pass

print('Done!')

# Destination and document title
doc_name = 'C:/User/<...>/Highlights_' + title + '.docx'

# Save document
doc.save(doc_name)

#Open the folder wher we save the document
os.startfile('C:/Users/<...>')

#Auto-Remove the first file "__.html", the file Kindle send to us when we export the highlights
os.remove(fichero)
